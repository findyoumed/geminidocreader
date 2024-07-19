# #7 DOCUMENTGPT
# DocumentGPT는 사용자가 업로드한 문서에 대해 질문하고 답변을 받을 수 있는 인공지능 기반 시스템입니다.

# #7.0 Introduction
# 이 섹션에서는 DocumentGPT에 필요한 기본 설정과 라이브러리를 소개합니다.

# 필요한 라이브러리들을 가져옵니다. 각 라이브러리는 특정 기능을 담당합니다.
import io  # 입출력 작업을 위한 도구입니다.
import os  # 운영 체제와 상호작용하기 위한 도구입니다.
import sys  # 파이썬 인터프리터와 관련된 기능을 제공합니다.
import tempfile  # 임시 파일을 생성하고 관리하는 도구입니다.
import traceback  # 오류 추적을 위한 도구입니다.
import streamlit as st  # 웹 애플리케이션을 쉽게 만들 수 있는 프레임워크입니다.

# 추가적인 기능을 위한 라이브러리들을 가져옵니다.
from dotenv import load_dotenv  # 환경 변수를 관리하는 도구입니다.
from langchain_community.document_loaders import PyPDFLoader  # PDF 파일을 읽기 위한 도구입니다.
from langchain_community.vectorstores.faiss import FAISS  # 벡터 데이터베이스를 사용하기 위한 도구입니다.
from langchain_huggingface import HuggingFaceEmbeddings  # 텍스트를 벡터로 변환하는 도구입니다.
from langchain_core.prompts import ChatPromptTemplate  # AI와의 대화를 위한 템플릿을 만드는 도구입니다.
from langchain_core.output_parsers import StrOutputParser  # AI의 출력을 처리하는 도구입니다.
from langchain_core.runnables import RunnablePassthrough, RunnableLambda  # 여러 작업을 연결하는 도구입니다.
from langchain_core.callbacks import BaseCallbackHandler  # AI의 응답 과정을 처리하는 도구입니다.
from langchain.globals import set_verbose  # 상세한 로그를 출력하도록 설정하는 도구입니다.
from langchain_google_genai import GoogleGenerativeAI  # Google의 AI 모델을 사용하기 위한 도구입니다.
from docx import Document  # Word 문서를 읽기 위한 도구입니다.
from langchain_core.documents import Document as LangchainDocument  # Langchain에서 사용하는 문서 형식입니다.

# 환경 변수를 로드합니다. 이는 API 키 등의 민감한 정보를 안전하게 관리하기 위함입니다.
load_dotenv()

# Streamlit 앱의 페이지 설정을 합니다. 제목과 아이콘을 지정합니다.
st.set_page_config(page_title="DocumentGPT", page_icon="📃")

# Langchain의 상세 로깅을 활성화합니다. 이는 디버깅에 유용합니다.
set_verbose(True)

# #7.1 Magic
# 이 섹션에서는 파일 처리와 벡터화 과정을 다룹니다.

# 예외 처리를 위한 함수를 정의합니다. 오류가 발생했을 때 상세 정보를 출력합니다.
def handle_exception(exc_type, exc_value, exc_traceback):
    # 오류 메시지를 표준 에러 스트림에 출력합니다.
    print("An error occurred:", file=sys.stderr)
    # 상세한 오류 정보를 출력합니다.
    traceback.print_exception(exc_type, exc_value, exc_traceback, file=sys.stderr)

# 정의한 예외 처리 함수를 시스템의 예외 처리기로 설정합니다.
sys.excepthook = handle_exception

# AI의 응답을 실시간으로 처리하기 위한 콜백 클래스를 정의합니다.
class ChatCallbackHandler(BaseCallbackHandler):
    message = ""  # AI의 응답을 저장할 변수입니다.

    def on_llm_start(self, *args, **kwargs):
        # AI가 응답을 시작할 때 호출됩니다.
        self.message_box = st.empty()  # Streamlit에 빈 컨테이너를 생성합니다.

    def on_llm_end(self, *args, **kwargs):
        # AI가 응답을 완료했을 때 호출됩니다.
        save_message(self.message, "ai")  # 완성된 응답을 저장합니다.

    def on_llm_new_token(self, token, *args, **kwargs):
        # AI가 새로운 토큰을 생성할 때마다 호출됩니다.
        self.message += token  # 새 토큰을 기존 메시지에 추가합니다.
        self.message_box.markdown(self.message)  # 업데이트된 메시지를 실시간으로 표시합니다.

# Google의 AI 모델을 초기화합니다.
llm = GoogleGenerativeAI(
    model="gemini-pro",  # 사용할 모델 이름입니다.
    temperature=0.1,  # 응답의 창의성을 조절하는 값입니다. 낮을수록 더 일관된 응답을 생성합니다.
    google_api_key=os.getenv("GOOGLE_API_KEY"),  # Google API 키를 환경 변수에서 가져옵니다.
)

# PDF 파일을 처리하는 함수입니다.
def process_pdf(file):
    # 임시 파일을 생성하여 업로드된 PDF 내용을 저장합니다.
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
        temp_file.write(file.getvalue())
        temp_file_path = temp_file.name

    # PDF 로더를 사용해 파일 내용을 읽어옵니다.
    loader = PyPDFLoader(temp_file_path)
    docs = loader.load()

    # 사용이 끝난 임시 파일을 삭제합니다.
    os.unlink(temp_file_path)

    return docs  # 처리된 PDF 내용을 반환합니다.

# Word 문서를 처리하는 함수입니다.
def process_docx(file):
    # 파일 내용을 바이트 스트림으로 읽어옵니다.
    docx_content = io.BytesIO(file.getvalue())
    # python-docx 라이브러리를 사용해 Word 문서를 엽니다.
    doc = Document(docx_content)
    # 문서의 모든 단락을 하나의 문자열로 합칩니다.
    text_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    # Langchain 문서 형식으로 변환하여 반환합니다.
    return [
        LangchainDocument(page_content=text_content, metadata={"source": file.name})
    ]

# #7.2 Data Flow
# 이 섹션에서는 데이터의 흐름과 처리 과정을 설명합니다.

# 파일을 벡터화하는 함수입니다. @st.cache_data 데코레이터는 결과를 캐시에 저장하여 재사용합니다.
@st.cache_data(show_spinner="Processing file...")
def embed_file(file):
    try:
        # 파일의 확장자를 확인합니다.
        file_extension = file.name.split(".")[-1].lower()

        # 파일 확장자에 따라 적절한 처리 방법을 선택합니다.
        if file_extension == "pdf":
            docs = process_pdf(file)
        elif file_extension == "docx":
            docs = process_docx(file)
        elif file_extension == "txt":
            # 텍스트 파일의 내용을 읽어 Langchain 문서 형식으로 변환합니다.
            content = file.getvalue().decode("utf-8", errors="ignore")
            docs = [
                LangchainDocument(page_content=content, metadata={"source": file.name})
            ]
        else:
            # 지원하지 않는 파일 형식일 경우 에러 메시지를 표시합니다.
            st.error(f"Unsupported file type: {file_extension}")
            return None

        # 문서 내용을 벡터로 변환하는 모델을 초기화합니다.
        embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
        # 문서 벡터를 FAISS에 저장합니다.
        vectorstore = FAISS.from_documents(docs, embeddings)
        # FAISS 벡터 저장소를 검색 가능한 형태로 변환합니다.
        retriever = vectorstore.as_retriever()

        return retriever  # 검색기를 반환합니다.
    except Exception as e:
        # 파일 처리 중 오류가 발생하면 에러 메시지를 표시합니다.
        st.error(f"An error occurred while processing the file: {e}")
        return None

# #7.3 Multi Page
# 이 섹션에서는 여러 페이지로 구성된 문서 처리 방법을 설명합니다.

# 메시지를 저장하는 함수입니다.
def save_message(message, role):
    # 세션 상태에 메시지 리스트가 없으면 새로 만듭니다.
    if "messages" not in st.session_state:
        st.session_state["messages"] = []
    # 새 메시지를 리스트에 추가합니다.
    st.session_state["messages"].append({"message": message, "role": role})

# 메시지를 화면에 표시하고 저장하는 함수입니다.
def send_message(message, role, save=True):
    # Streamlit의 chat_message 기능을 사용해 메시지를 표시합니다.
    with st.chat_message(role):
        st.markdown(message)
    # save 매개변수가 True면 메시지를 저장합니다.
    if save:
        save_message(message, role)

# 이전 대화 내용을 화면에 표시하는 함수입니다.
def paint_history():
    # 세션 상태에 저장된 모든 메시지를 화면에 표시합니다.
    if "messages" in st.session_state:
        for message in st.session_state["messages"]:
            send_message(message["message"], message["role"], save=False)

# 여러 문서의 내용을 하나의 문자열로 합치는 함수입니다.
def format_docs(docs):
    # 각 문서의 내용을 줄바꿈으로 구분하여 하나의 긴 문자열로 만듭니다.
    return "\n\n".join(doc.page_content for doc in docs)

# #7.4 Chat Messages
# 이 섹션에서는 AI와의 대화를 위한 프롬프트 설정을 다룹니다.

# AI 응답을 위한 프롬프트 템플릿을 설정합니다.
prompt = ChatPromptTemplate.from_messages(
    [
        (
            "system",
            """
    당신은 도움이 되는 AI 비서예요. 주어진 정보를 바탕으로 질문에 답하세요. 
    정보가 부족하면 일반적인 지식을 사용해도 돼요. 하지만 어떤 부분이 주어진 정보에서 나온 것이고 
    어떤 부분이 일반 지식인지 명확히 말해주세요.
    잘 모르겠거나 질문과 관계없는 내용이면 솔직히 그렇다고 말해주세요.
    
    주어진 정보: {context}
    """,
        ),
        ("human", "{question}"),
    ]
)

# #7.5 Recap
# 이 섹션에서는 지금까지 배운 내용을 요약합니다.

# Streamlit 앱의 메인 제목을 설정합니다.
st.title("DocumentGPT")

# 앱 사용법을 안내하는 메시지를 표시합니다.
st.markdown(
    """
안녕하세요! 이 앱은 여러분의 문서에 대해 질문하면 답변해주는 똑똑한 비서예요.
왼쪽에서 파일을 올려주세요, 그러면 시작할 수 있어요!
"""
)

# #7.6 Uploading Documents (계속)
# 이 섹션에서는 문서 업로드 기능을 구현합니다.

# Streamlit의 사이드바에 파일 업로드 버튼을 만듭니다.
with st.sidebar:
    uploaded_file = st.file_uploader(
        "PDF, TXT, DOCX 파일 중 하나를 올려주세요", type=["pdf", "txt", "docx"]
    )
    # 사용자가 PDF, TXT, DOCX 형식의 파일만 업로드할 수 있도록 제한합니다.

# #7.7 Chat History
# 이 섹션에서는 대화 기록을 관리하는 방법을 설명합니다.

# 파일이 업로드되면 이 부분이 실행됩니다.
if uploaded_file is not None:
    # 업로드된 파일을 처리하고 검색 가능한 형태로 변환합니다.
    retriever = embed_file(uploaded_file)
    
    # 파일 처리가 성공적으로 끝나면 이 부분이 실행됩니다.
    if retriever:
        # AI가 준비되었다는 메시지를 화면에 표시합니다.
        send_message("좋아요! 이제 질문해주세요!", "ai", save=False)
        
        # 이전 대화 내용이 있다면 화면에 표시합니다.
        paint_history()
        
        # 사용자가 질문을 입력할 수 있는 채팅 입력 창을 만듭니다.
        message = st.chat_input("파일에 대해 궁금한 점을 물어보세요...")
        
        # 사용자가 질문을 입력하면 이 부분이 실행됩니다.
        if message:
            # 사용자의 질문을 화면에 표시하고 저장합니다.
            send_message(message, "human")
            
            # #7.8 Chain
            # 이 섹션에서는 여러 처리 단계를 연결하는 체인(Chain)을 구현합니다.
            
            # 여러 단계의 처리 과정을 하나로 연결합니다:
            chain = (
                {
                    "context": retriever | RunnableLambda(format_docs),
                    # retriever로 관련 문서를 검색하고, format_docs로 문서 내용을 정리합니다.
                    "question": RunnablePassthrough(),
                    # 사용자의 질문(question)은 그대로 전달합니다.
                }
                | prompt
                # 정리된 문서 내용과 질문을 prompt 템플릿에 넣어 AI에게 전달할 최종 프롬프트를 만듭니다.
                | llm
                # llm(AI 모델)에 프롬프트를 전달하여 답변을 생성합니다.
                | StrOutputParser()
                # 생성된 답변을 문자열로 파싱합니다.
            )
            
            # #7.9 Streaming
            # 이 섹션에서는 AI의 응답을 실시간으로 스트리밍하는 방법을 구현합니다.
            
            # AI의 답변을 화면에 표시할 준비를 합니다.
            with st.chat_message("ai"):
                # chain을 실행하여 AI의 답변을 생성합니다.
                response = chain.invoke(message)
                # AI의 답변을 화면에 표시합니다.
                st.write(response)
                # AI의 답변을 저장합니다.
                save_message(response, "ai")
else:
    # 파일이 아직 업로드되지 않았다면 이 부분이 실행됩니다.
    # 이전 대화 내용을 모두 지웁니다. 새로운 대화를 시작할 준비를 하는 것입니다.
    st.session_state["messages"] = []

# #7.10 Recap
# 이 섹션에서는 전체 코드의 주요 기능을 다시 한 번 요약합니다:
# 1. 사용자가 문서를 업로드하면 해당 문서를 처리하고 벡터화합니다.
# 2. 사용자가 질문을 입력하면 관련된 문서 내용을 검색합니다.
# 3. 검색된 내용과 사용자의 질문을 바탕으로 AI에게 전달할 프롬프트를 생성합니다.
# 4. AI 모델이 프롬프트를 바탕으로 답변을 생성합니다.
# 5. 생성된 답변을 실시간으로 화면에 표시하고 저장합니다.
# 이 과정을 통해 사용자는 자신의 문서에 대해 질문하고 답변을 받을 수 있습니다.